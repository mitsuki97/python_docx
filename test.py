import xlrd
import matplotlib
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# 获取学习成绩信息
def GetExcelInfo():
    print("开始获取表格内容信息")

    # 打开指定文档
    xlsx = xlrd.open_workbook('score.xlsx')

    # 获取sheet
    sheet = xlsx.sheet_by_index(0)

    # 获取表格行数
    nrows = sheet.nrows
    print("一共 ",nrows," 行数据")

    # 获取第2列，和第4列 所有值（列表生成式）,从第2行开始获取
    nameList = [str(sheet.cell_value(i, 1)) for i in range(1, nrows)]
    scoreList = [int(sheet.cell_value(i, 3)) for i in range(1, nrows)]

    # 返回名字列表和分数列表
    return nameList,scoreList

# 生成学生成绩柱状图（使用matplotlib）
# 会生成一张名为"studentScore.jpg"的图片
def GenerateScorePic(scoreList):
    # 解析成绩列表，生成横纵坐标列表
    xNameList = [str(studentInfo[0]) for studentInfo in scoreList]
    yScoreList = [int(studentInfo[1]) for studentInfo in scoreList]
    print("xNameList",xNameList)
    print("yScoreList",yScoreList)

    # 设置字体格式
    matplotlib.rcParams['font.sans-serif'] = ['SimHei']  # 用黑体显示中文

    # 设置绘图尺寸
    plt.figure(figsize=(10,5))

    # 绘制图像
    plt.bar(x=xNameList, height=yScoreList, label='学生成绩', color='steelblue', alpha=0.8)

    # 在柱状图上显示具体数值, ha参数控制水平对齐方式, va控制垂直对齐方式
    for x1, yy in scoreList:
        plt.text(x1, yy + 1, str(yy), ha='center', va='bottom', fontsize=16, rotation=0)

    # 设置标题
    plt.title("学生成绩柱状图")

    # 为两条坐标轴设置名称
    plt.xlabel("学生姓名")
    plt.ylabel("学生成绩")

    # 显示图例
    plt.legend()

    # 坐标轴旋转
    plt.xticks(rotation=90)

    # 设置底部比例，防止横坐标显示不全
    plt.gcf().subplots_adjust(bottom=0.25)

    # 保存为图片
    plt.savefig("studentScore.jpg")

    # 直接显示
    plt.show()

# 开始生成报告
def GenerateScoreReport(scoreOrder,picPath):
    # 新建一个文档
    document = Document()

    # 设置标题
    document.add_heading('数据分析报告', 0)

    # 添加第一名的信息
    p1 = document.add_paragraph("分数排在第一的学生姓名为: ")
    p1.add_run(scoreOrder[0][0]).bold = True
    p1.add_run(" 分数为: ")
    p1.add_run(str(scoreOrder[0][1])).italic = True

    # 添加总体情况信息
    p2 = document.add_paragraph("共有: ")
    p2.add_run(str(len(scoreOrder))).bold = True
    p2.add_run(" 名学生参加了考试，学生考试的总体情况: ")

    # 添加考试情况表格
    table = document.add_table(rows=1, cols=2)
    table.style = 'Medium Grid 1 Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '学生姓名'
    hdr_cells[1].text = '学生分数'
    for studentName,studentScore in scoreOrder:
        row_cells = table.add_row().cells
        row_cells[0].text = studentName
        row_cells[1].text = str(studentScore)

    # 添加学生成绩柱状图
    document.add_picture(picPath, width=Inches(6))

    document.save('学生成绩报告.docx')

if __name__ == "__main__":
    # 调用信息获取方法，获取用户信息
    nameList,scoreList = GetExcelInfo()
    # print("nameList:",nameList)
    # print("ScoreList:",scoreList)

    # 将名字和分数列表合并成字典(将学生姓名和分数关联起来)
    scoreDictionary = dict(zip(nameList, scoreList))
    # print("dictionary:",scoreDictionary)

    # 对字典进行值排序，高分在前,reverse=True 代表降序排列
    scoreOrder = sorted(scoreDictionary.items(), key=lambda x: x[1], reverse=True)
    # print("scoreOrder",scoreOrder)

    # 将进行排序后的学生成绩列表生成柱状图
    GenerateScorePic(scoreOrder)

    # 开始生成报告
    picPath = "studentScore.jpg"
    GenerateScoreReport(scoreOrder,picPath)

    print("任务完成，报表生成完毕！")
